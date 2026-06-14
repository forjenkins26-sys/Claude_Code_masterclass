"""
QA Automation Portfolio Word Document Generator
Generates impressive .docx with architecture diagrams, color sections, tables
"""

import io
import math
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import PIL.Image
import PIL.ImageDraw
import PIL.ImageFont

# ─── Color Palette ────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x0D, 0x1B, 0x2A)   # deep navy — headings
TEAL       = RGBColor(0x00, 0x87, 0x8A)   # teal — accent / highlights
GREEN      = RGBColor(0x27, 0xAE, 0x60)   # pass / success
RED        = RGBColor(0xC0, 0x39, 0x2B)   # fail / blocked
ORANGE     = RGBColor(0xE6, 0x7E, 0x22)   # warning / in progress
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF4, 0xF6, 0xF9)
MID_GRAY   = RGBColor(0x95, 0xA5, 0xA6)
DARK_GRAY  = RGBColor(0x2C, 0x3E, 0x50)
GOLD       = RGBColor(0xF3, 0x9C, 0x12)

# ─── Helpers ──────────────────────────────────────────────────────────────────

def rgb_to_hex(rgb: RGBColor) -> str:
    # RGBColor is a tuple subclass: (r, g, b)
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = rgb_to_hex(rgb)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(table, border_color="AAAAAA"):
    """Add borders to all cells in a table."""
    tbl = table._tbl
    for row in tbl.findall(qn('w:tr')):
        for cell in row.findall(qn('w:tc')):
            tcPr = cell.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), border_color)
                tcBorders.append(border)
            tcPr.append(tcBorders)

def add_paragraph_spacing(doc, before=0, after=6):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    return p

def add_heading(doc, text, level=1, color=NAVY, size=22, bold=True, after=6):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12 if level == 1 else 8)
    pf.space_after = Pt(after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if level == 1:
        # Add underline accent
        run.font.color.rgb = NAVY
    return p

def add_body(doc, text, size=10.5, color=DARK_GRAY, bold=False, italic=False, before=2, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text, size=10, indent=0.3):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK_GRAY
    return p

def add_code_block(doc, code_text):
    """Add styled code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Shade paragraph background via shading on the paragraph itself
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1E1E2E')
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0xA8, 0xFF, 0x78)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '008789')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ─── Image Generators ─────────────────────────────────────────────────────────

def make_architecture_diagram():
    """8-layer architecture pyramid diagram."""
    W, H = 900, 520
    img = PIL.Image.new('RGB', (W, H), (248, 250, 252))
    draw = PIL.ImageDraw.Draw(img)

    layers = [
        ("8  TESTS",            (13, 71, 161),  (227, 242, 253)),
        ("7  SERVICES",         (1, 87, 155),   (225, 245, 254)),
        ("6  PAGES (POM)",      (0, 131, 143),  (224, 247, 250)),
        ("5  COMPONENTS",       (0, 121, 107),  (224, 242, 241)),
        ("4  API CLIENT",       (46, 125, 50),  (232, 245, 233)),
        ("3  TEST DATA",        (130, 119, 23), (255, 253, 231)),
        ("2  UTILITIES",        (183, 84, 9),   (255, 243, 224)),
        ("1  CONFIG / ENV",     (136, 14, 79),  (252, 228, 236)),
    ]

    bar_h = 50
    margin_x = 60
    gap = 6
    total = len(layers)

    for i, (label, dark, light) in enumerate(layers):
        y = 20 + i * (bar_h + gap)
        shrink = i * 18
        x0 = margin_x + shrink
        x1 = W - margin_x - shrink
        # Shadow
        draw.rectangle([x0+3, y+3, x1+3, y+bar_h+3], fill=(200, 200, 200))
        # Main bar
        draw.rectangle([x0, y, x1, y+bar_h], fill=light, outline=dark, width=2)
        # Label
        try:
            font = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 17)
        except:
            font = PIL.ImageFont.load_default()
        tw = draw.textlength(label, font=font)
        draw.text(((x0+x1)//2 - tw//2, y + bar_h//2 - 9), label, fill=dark, font=font)

    # Title
    try:
        title_font = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 20)
    except:
        title_font = PIL.ImageFont.load_default()
    draw.text((W//2 - 160, H-38), "8-Layer POM Architecture", fill=(13, 27, 42), font=title_font)

    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def make_pipeline_diagram():
    """CI/QA pipeline flow diagram."""
    W, H = 960, 200
    img = PIL.Image.new('RGB', (W, H), (248, 250, 252))
    draw = PIL.ImageDraw.Draw(img)

    steps = [
        ("Jira Epic\n(MCP fetch)", (13, 71, 161), (227, 242, 253)),
        ("/test-case\ncreation", (0, 131, 143), (224, 247, 250)),
        ("Playwright\nExecution", (46, 125, 50), (232, 245, 233)),
        ("Allure\nReport", (130, 75, 14), (255, 243, 224)),
        ("Jira Status\nUpdate", (136, 14, 79), (252, 228, 236)),
    ]

    box_w = 140
    box_h = 80
    gap = 28
    total_w = len(steps) * box_w + (len(steps)-1) * gap
    start_x = (W - total_w) // 2
    y = (H - box_h) // 2

    try:
        font = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 13)
        small = PIL.ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 11)
    except:
        font = PIL.ImageFont.load_default()
        small = font

    for i, (label, dark, light) in enumerate(steps):
        x0 = start_x + i * (box_w + gap)
        x1 = x0 + box_w
        # Shadow
        draw.rounded_rectangle([x0+3, y+3, x1+3, y+box_h+3], radius=10, fill=(200,200,200))
        # Box
        draw.rounded_rectangle([x0, y, x1, y+box_h], radius=10, fill=light, outline=dark, width=2)
        # Step number badge
        draw.ellipse([x0+6, y+6, x0+24, y+24], fill=dark)
        try:
            num_font = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 11)
        except:
            num_font = font
        draw.text((x0+11, y+7), str(i+1), fill=(255,255,255), font=num_font)
        # Label (multiline)
        lines = label.split('\n')
        text_y = y + box_h//2 - len(lines)*8
        for line in lines:
            tw = draw.textlength(line, font=font)
            draw.text(((x0+x1)//2 - tw//2, text_y), line, fill=dark, font=font)
            text_y += 16

        # Arrow between boxes
        if i < len(steps) - 1:
            ax = x1 + 2
            ay = y + box_h // 2
            draw.line([(ax, ay), (ax + gap - 4, ay)], fill=(13, 71, 161), width=3)
            # Arrowhead
            draw.polygon([(ax+gap-4, ay-6), (ax+gap+4, ay), (ax+gap-4, ay+6)], fill=(13,71,161))

    # Title
    try:
        title_font = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 16)
    except:
        title_font = font
    draw.text((W//2 - 160, 10), "End-to-End QA Automation Pipeline", fill=(13, 27, 42), font=title_font)

    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def make_jira_lifecycle():
    """Jira status lifecycle diagram."""
    W, H = 820, 160
    img = PIL.Image.new('RGB', (W, H), (248, 250, 252))
    draw = PIL.ImageDraw.Draw(img)

    states = [
        ("TO DO",       (120, 144, 156), (236, 239, 241)),
        ("IN PROGRESS", (25, 118, 210),  (227, 242, 253)),
        ("DONE",        (56, 142, 60),   (232, 245, 233)),
        ("BLOCKED",     (198, 40, 40),   (255, 235, 238)),
    ]

    box_w = 140
    box_h = 56
    gap = 36
    total_w = len(states) * box_w + (len(states)-1) * gap
    start_x = (W - total_w) // 2
    y = (H - box_h) // 2 + 10

    try:
        font = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 13)
    except:
        font = PIL.ImageFont.load_default()

    for i, (label, dark, light) in enumerate(states):
        x0 = start_x + i * (box_w + gap)
        x1 = x0 + box_w
        draw.rounded_rectangle([x0+2, y+2, x1+2, y+box_h+2], radius=8, fill=(180,180,180))
        draw.rounded_rectangle([x0, y, x1, y+box_h], radius=8, fill=light, outline=dark, width=2)
        tw = draw.textlength(label, font=font)
        draw.text(((x0+x1)//2 - tw//2, y + box_h//2 - 9), label, fill=dark, font=font)

        if i < len(states) - 1:
            # Fork arrow after IN PROGRESS → DONE and BLOCKED
            ax = x1 + 2
            ay = y + box_h // 2
            if i == 1:
                # Arrow to DONE (top fork)
                draw.line([(ax, ay), (ax+gap-4, ay)], fill=(56,142,60), width=2)
                draw.polygon([(ax+gap-4, ay-5), (ax+gap+4, ay), (ax+gap-4, ay+5)], fill=(56,142,60))
            else:
                draw.line([(ax, ay), (ax+gap-4, ay)], fill=(100,100,100), width=2)
                draw.polygon([(ax+gap-4, ay-5), (ax+gap+4, ay), (ax+gap-4, ay+5)], fill=(100,100,100))

    # Label for fork
    try:
        small = PIL.ImageFont.truetype("C:/Windows/Fonts/ariali.ttf", 10)
    except:
        small = font
    draw.text((start_x + 2*(box_w+gap) + 4, y - 16), "pass → Done | app bug → Blocked", fill=(80,80,80), font=small)

    try:
        title_font = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 15)
    except:
        title_font = font
    draw.text((W//2 - 130, 8), "Jira Test Case Lifecycle (Auto-Managed)", fill=(13,27,42), font=title_font)

    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def make_test_metrics():
    """Test results donut + bar chart."""
    W, H = 860, 280
    img = PIL.Image.new('RGB', (W, H), (248, 250, 252))
    draw = PIL.ImageDraw.Draw(img)

    # ── Donut chart (left) ──
    cx, cy, r_outer, r_inner = 170, 140, 110, 55
    data = [
        ("Pass",    101, (39, 174, 96)),
        ("Blocked",   2, (192, 57, 43)),
    ]
    total = sum(d[1] for d in data)
    start = -90
    for label, val, color in data:
        sweep = 360 * val / total
        end = start + sweep
        draw.pieslice(
            [cx-r_outer, cy-r_outer, cx+r_outer, cy+r_outer],
            start=start, end=end, fill=color
        )
        start = end
    # White circle to make donut
    draw.ellipse([cx-r_inner, cy-r_inner, cx+r_inner, cy+r_inner], fill=(248,250,252))
    # Center text
    try:
        big = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 22)
        sm  = PIL.ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 12)
    except:
        big = PIL.ImageFont.load_default()
        sm = big
    draw.text((cx-18, cy-18), "103", fill=(13,27,42), font=big)
    draw.text((cx-14, cy+4), "Tests", fill=(80,80,80), font=sm)

    # Legend
    lx, ly = 290, 80
    for label, val, color in data:
        draw.rectangle([lx, ly, lx+16, ly+16], fill=color)
        draw.text((lx+22, ly), f"{label}: {val} ({100*val//total}%)", fill=(44,62,80), font=sm)
        ly += 28

    # ── Bar chart (right) — tests per epic ──
    epics = [
        ("Login\n(TC)", 17, (13,71,161)),
        ("Reg\n(REG)", 19, (0,131,143)),
        ("Forgot\n(FP)", 15, (46,125,50)),
        ("Blinkit\n(BL)", 19, (130,75,14)),
        ("Reg-Demo\n(REG)", 33, (136,14,79)),
    ]
    max_val = max(e[1] for e in epics)
    bx_start = 430
    bar_w = 70
    bar_gap = 18
    chart_h = 160
    chart_base_y = 220

    try:
        bar_font = PIL.ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 10)
        val_font = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 11)
    except:
        bar_font = PIL.ImageFont.load_default()
        val_font = bar_font

    for i, (label, val, color) in enumerate(epics):
        x0 = bx_start + i * (bar_w + bar_gap)
        bar_height = int(chart_h * val / max_val)
        y0 = chart_base_y - bar_height
        # Shadow
        draw.rectangle([x0+3, y0+3, x0+bar_w+3, chart_base_y+3], fill=(200,200,200))
        # Bar
        draw.rectangle([x0, y0, x0+bar_w, chart_base_y], fill=color)
        # Value on top
        vw = draw.textlength(str(val), font=val_font)
        draw.text((x0 + bar_w//2 - vw//2, y0-16), str(val), fill=color, font=val_font)
        # Label below (multiline)
        lines = label.split('\n')
        ty = chart_base_y + 6
        for line in lines:
            lw = draw.textlength(line, font=bar_font)
            draw.text((x0 + bar_w//2 - lw//2, ty), line, fill=(80,80,80), font=bar_font)
            ty += 13

    # Titles
    try:
        title_font = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 14)
    except:
        title_font = PIL.ImageFont.load_default()
    draw.text((80, 12), "Test Results Overview", fill=(13,27,42), font=title_font)
    draw.text((565, 12), "Tests per Epic", fill=(13,27,42), font=title_font)

    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def make_ah_rules_diagram():
    """Anti-Hallucination two-source model diagram."""
    W, H = 860, 240
    img = PIL.Image.new('RGB', (W, H), (248, 250, 252))
    draw = PIL.ImageDraw.Draw(img)

    try:
        bold = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 14)
        norm = PIL.ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 11)
        sm   = PIL.ImageFont.truetype("C:/Windows/Fonts/ariali.ttf", 10)
    except:
        bold = norm = sm = PIL.ImageFont.load_default()

    # Left box — Jira Epic (requirements)
    draw.rounded_rectangle([30, 60, 280, 190], radius=12, fill=(227,242,253), outline=(13,71,161), width=2)
    draw.text((55, 72), "Jira Epic / Requirements", fill=(13,71,161), font=bold)
    for i, line in enumerate(["Acceptance criteria", "Expected error text", "Navigation targets", "Success states"]):
        draw.text((50, 100 + i*20), f"  ✓  {line}", fill=(25,60,130), font=norm)

    # Right box — UI DOM
    draw.rounded_rectangle([580, 60, 830, 190], radius=12, fill=(232,245,233), outline=(46,125,50), width=2)
    draw.text((610, 72), "UI Analysis (DOM / WebFetch)", fill=(46,125,50), font=bold)
    for i, line in enumerate(["Element IDs / selectors", "Input types + attributes", "Button text (locator only)", "Form structure"]):
        draw.text((600, 100 + i*20), f"  ✓  {line}", fill=(27,94,32), font=norm)

    # Center — Test Case
    draw.rounded_rectangle([320, 80, 540, 170], radius=14, fill=(13,71,161), outline=(0,0,0), width=1)
    draw.text((345, 96), "TEST CASE", fill=(255,255,255), font=bold)
    draw.text((330, 122), "Assertion  ←  Epic", fill=(255,220,100), font=norm)
    draw.text((330, 142), "Locator   ←  DOM", fill=(180,255,180), font=norm)

    # Arrows
    draw.line([(280, 125), (320, 125)], fill=(13,71,161), width=3)
    draw.polygon([(316,119),(326,125),(316,131)], fill=(13,71,161))
    draw.line([(540, 125), (580, 125)], fill=(46,125,50), width=3)
    draw.polygon([(540,119),(540,131),(551,125)], fill=(46,125,50))

    # Title
    try:
        title_font = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 16)
    except:
        title_font = bold
    draw.text((W//2-200, 14), "Anti-Hallucination Rule 19 — Two-Source Model", fill=(13,27,42), font=title_font)

    # Bottom note
    draw.text((100, 205), "❌ Wrong: derive expected result from UI observation → validates broken implementation",
              fill=(192,57,43), font=sm)
    draw.text((100, 220), "✅ Right:  derive expected result from Epic requirement → catches broken implementation",
              fill=(39,174,96), font=sm)

    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def make_autofix_diagram():
    """Auto-Fix Protocol flowchart."""
    W, H = 860, 300
    img = PIL.Image.new('RGB', (W, H), (248, 250, 252))
    draw = PIL.ImageDraw.Draw(img)

    try:
        bold = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 13)
        norm = PIL.ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 11)
    except:
        bold = norm = PIL.ImageFont.load_default()

    steps = [
        ("DETECT",      "Extract what/\nwhere/why failed",     (183,28,28),  (255,235,238)),
        ("INVESTIGATE", "Headed mode +\nDOM inspection",        (25,118,210), (227,242,253)),
        ("CLASSIFY",    "Test bug or\nApp bug?",               (123,31,162), (243,229,245)),
        ("FIX / FILE",  "Fix test OR\nFile Jira bug",          (0,131,143),  (224,247,250)),
        ("VERIFY 3x",   "Single→Related\n→Full suite",         (46,125,50),  (232,245,233)),
    ]

    box_w = 130
    box_h = 80
    gap = 20
    total_w = len(steps) * box_w + (len(steps)-1) * gap
    start_x = (W - total_w) // 2
    y = 80

    for i, (title, subtitle, dark, light) in enumerate(steps):
        x0 = start_x + i * (box_w + gap)
        x1 = x0 + box_w
        # Shadow
        draw.rounded_rectangle([x0+3,y+3,x1+3,y+box_h+3], radius=10, fill=(180,180,180))
        # Box
        draw.rounded_rectangle([x0,y,x1,y+box_h], radius=10, fill=light, outline=dark, width=2)
        # Step badge
        draw.ellipse([x0+6, y+6, x0+24, y+24], fill=dark)
        draw.text((x0+10, y+7), str(i+1), fill=(255,255,255), font=norm)
        # Title
        tw = draw.textlength(title, font=bold)
        draw.text(((x0+x1)//2 - tw//2, y+22), title, fill=dark, font=bold)
        # Subtitle
        for j, line in enumerate(subtitle.split('\n')):
            lw = draw.textlength(line, font=norm)
            draw.text(((x0+x1)//2 - lw//2, y+44+j*14), line, fill=(60,60,60), font=norm)

        # Arrow
        if i < len(steps)-1:
            ax = x1+2
            ay = y + box_h//2
            draw.line([(ax, ay),(ax+gap-4, ay)], fill=dark, width=2)
            draw.polygon([(ax+gap-5,ay-5),(ax+gap+3,ay),(ax+gap-5,ay+5)], fill=dark)

    # Bottom label
    draw.text((W//2-160, y+box_h+20),
              "Max 3 attempts. Escalate only if stuck. Never ask user before attempting fix.",
              fill=(80,80,80), font=norm)

    # Rule labels below boxes
    labels = ["AFP-1,6", "AFP-7,9,14", "AFP-13,15", "AFP-10,13", "AFP-4,11"]
    for i, lbl in enumerate(labels):
        x0 = start_x + i*(box_w+gap)
        lw = draw.textlength(lbl, font=norm)
        draw.text((x0+box_w//2-lw//2, y+box_h+6), lbl, fill=(120,120,120), font=norm)

    # Title
    try:
        title_font = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 16)
    except:
        title_font = bold
    draw.text((W//2-140, 20), "Auto-Fix Protocol — 5-Step Process", fill=(13,27,42), font=title_font)

    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf

def make_accuracy_chart():
    """Side-by-side horizontal bar chart: accuracy with vs without frameworks."""
    W, H = 900, 420
    img = PIL.Image.new('RGB', (W, H), (248, 250, 252))
    draw = PIL.ImageDraw.Draw(img)

    try:
        bold  = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 13)
        norm  = PIL.ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 11)
        small = PIL.ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 10)
        big   = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 16)
    except:
        bold = norm = small = big = PIL.ImageFont.load_default()

    categories = [
        ("SPA Selector Accuracy",       31, 97),
        ("Variable Naming Accuracy",     40, 98),
        ("Test Assertion Correctness",   48, 97),
        ("Bug Classification",           35, 95),
        ("Constraint Test Accuracy",      0, 98),
        ("URL/Navigation Accuracy",      30, 97),
        ("Jira Status Accuracy",         65, 99),
    ]

    label_w  = 230
    bar_area = W - label_w - 120
    max_val  = 100
    row_h    = 46
    start_y  = 55
    bar_gap  = 4

    # Title
    draw.text((W//2 - 220, 12), "Accuracy: Without vs With Anti-Hallucination + Auto-Fix Frameworks",
              fill=(13, 27, 42), font=big)

    # Legend
    draw.rectangle([label_w, 34, label_w+16, 48], fill=(192, 57, 43))
    draw.text((label_w + 20, 34), "Without Frameworks", fill=(80, 80, 80), font=small)
    draw.rectangle([label_w + 170, 34, label_w + 186, 48], fill=(39, 174, 96))
    draw.text((label_w + 190, 34), "With Frameworks", fill=(80, 80, 80), font=small)

    for i, (label, without, with_fw) in enumerate(categories):
        y = start_y + i * row_h
        # Row background
        bg = (240, 242, 245) if i % 2 == 0 else (248, 250, 252)
        draw.rectangle([0, y, W, y + row_h - 2], fill=bg)

        # Category label
        draw.text((10, y + row_h//2 - 8), label, fill=(44, 62, 80), font=norm)

        # Without bar (red)
        bar_w_without = int(bar_area * without / max_val)
        bx = label_w
        by_top = y + 4
        by_bot = y + row_h//2 - bar_gap
        draw.rectangle([bx, by_top, bx + bar_w_without, by_bot], fill=(192, 57, 43))
        # Label
        draw.text((bx + bar_w_without + 4, by_top), f"{without}%", fill=(192, 57, 43), font=small)

        # With bar (green)
        bar_w_with = int(bar_area * with_fw / max_val)
        by_top2 = y + row_h//2 + bar_gap
        by_bot2 = y + row_h - 6
        draw.rectangle([bx, by_top2, bx + bar_w_with, by_bot2], fill=(39, 174, 96))
        draw.text((bx + bar_w_with + 4, by_top2), f"{with_fw}%", fill=(39, 174, 96), font=small)

        # Delta badge
        delta = with_fw - without
        dx = W - 72
        draw.rounded_rectangle([dx, y + 8, dx + 58, y + row_h - 10], radius=6,
                                fill=(13, 71, 161))
        dw = draw.textlength(f"+{delta}pp", font=bold)
        draw.text((dx + (58 - dw)//2, y + row_h//2 - 10), f"+{delta}pp",
                  fill=(255, 255, 255), font=bold)

    # Danger zone annotation for constraint tests
    draw.text((label_w + 2, start_y + 4 * row_h + 6),
              "fill() = 0% accuracy (silent false pass always)", fill=(192, 57, 43), font=small)

    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def make_correction_cycle_chart():
    """Stacked time comparison — correction cycles with vs without frameworks."""
    W, H = 900, 300
    img = PIL.Image.new('RGB', (W, H), (248, 250, 252))
    draw = PIL.ImageDraw.Draw(img)

    try:
        bold  = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 13)
        norm  = PIL.ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 11)
        small = PIL.ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 10)
        big   = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 15)
    except:
        bold = norm = small = big = PIL.ImageFont.load_default()

    # Title
    draw.text((W//2 - 270, 10), "AI-Assisted QA — Time Per Epic: Correction Cycles vs Actual Work",
              fill=(13, 27, 42), font=big)

    # Two stacked bar configs: WITHOUT and WITH
    bar_w = 220
    bar_max_h = 190
    base_y = 260

    # WITHOUT: ~7-15 hrs total → show as 100% = 14 hours
    # Segments: generation (45min), selector fixes (90min), assertion rewrite (8hr), jira (2hr)
    without_segments = [
        ("Generation\n45 min",    45,   (99, 149, 184)),
        ("Selector fixes\n90 min", 90,  (230, 126, 34)),
        ("Assertion rewrite\n8 hr", 480, (192, 57, 43)),
        ("Jira updates\n2 hr",    120,  (142, 68, 173)),
    ]
    # WITH: ~25 min total
    with_segments = [
        ("Generation\n25 min",  25, (99, 149, 184)),
        ("Corrections\n<30 min", 25, (39, 174, 96)),
    ]

    total_without = sum(s[1] for s in without_segments)  # 735 min
    total_with    = sum(s[1] for s in with_segments)      # 50 min
    scale = bar_max_h / total_without  # scale both to same axis

    # Draw WITHOUT bar (left)
    bx_left = 160
    y = base_y
    for label, mins, color in reversed(without_segments):
        h_px = max(int(mins * scale), 14)
        y -= h_px
        draw.rectangle([bx_left, y, bx_left + bar_w, y + h_px], fill=color,
                        outline=(255,255,255), width=1)
        # Label inside if tall enough
        lines = label.split('\n')
        if h_px > 22:
            ty = y + h_px//2 - len(lines)*7
            for line in lines:
                lw = draw.textlength(line, font=small)
                draw.text((bx_left + bar_w//2 - lw//2, ty), line, fill=(255,255,255), font=small)
                ty += 14

    # Total label above
    tw = draw.textlength("~7–15 hours", font=bold)
    draw.text((bx_left + bar_w//2 - tw//2, y - 22), "~7–15 hours", fill=(192,57,43), font=bold)

    # Draw WITH bar (right)
    bx_right = 520
    y2 = base_y
    for label, mins, color in reversed(with_segments):
        h_px = max(int(mins * scale), 14)
        y2 -= h_px
        draw.rectangle([bx_right, y2, bx_right + bar_w, y2 + h_px], fill=color,
                        outline=(255,255,255), width=1)
        lines = label.split('\n')
        if h_px > 22:
            ty = y2 + h_px//2 - len(lines)*7
            for line in lines:
                lw = draw.textlength(line, font=small)
                draw.text((bx_right + bar_w//2 - lw//2, ty), line, fill=(255,255,255), font=small)
                ty += 14

    # The gap visually — tiny green bar
    tw2 = draw.textlength("~20–30 min", font=bold)
    draw.text((bx_right + bar_w//2 - tw2//2, y2 - 22), "~20–30 min", fill=(39,174,96), font=bold)

    # X-axis labels
    xleft  = draw.textlength("AI WITHOUT Frameworks", font=bold)
    xright = draw.textlength("AI WITH Frameworks", font=bold)
    draw.text((bx_left  + bar_w//2 - xleft//2,  base_y + 6), "AI WITHOUT Frameworks", fill=(192,57,43), font=bold)
    draw.text((bx_right + bar_w//2 - xright//2, base_y + 6), "AI WITH Frameworks",    fill=(39,174,96), font=bold)

    # Arrow + "20x faster"
    mid_x = (bx_left + bar_w + bx_right) // 2
    draw.line([(bx_left + bar_w + 10, base_y - bar_max_h//2),
               (bx_right - 10, base_y - bar_max_h//2)], fill=(13,71,161), width=2)
    draw.polygon([(bx_right - 10, base_y - bar_max_h//2 - 5),
                  (bx_right - 2,  base_y - bar_max_h//2),
                  (bx_right - 10, base_y - bar_max_h//2 + 5)], fill=(13,71,161))
    label_20x = "~20-30x faster"
    lw20 = draw.textlength(label_20x, font=bold)
    draw.text((mid_x - lw20//2, base_y - bar_max_h//2 - 20), label_20x, fill=(13,71,161), font=bold)

    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def make_ricepot_diagram():
    """RICEPOT framework visual — 7 colored hexagonal/pill badges in a row."""
    W, H = 900, 180
    img = PIL.Image.new('RGB', (W, H), (248, 250, 252))
    draw = PIL.ImageDraw.Draw(img)

    letters = [
        ("R", "Role",         (192, 57,  43),  (255, 205, 200)),
        ("I", "Instructions", (230, 126,  34),  (255, 235, 205)),
        ("C", "Context",      (243, 156,  18),  (255, 248, 210)),
        ("E", "Example",      (39,  174,  96),  (210, 245, 220)),
        ("P", "Parameters",   (0,   135, 138),  (200, 242, 244)),
        ("O", "Output",       (41,  128, 185),  (210, 235, 255)),
        ("T", "Tone",         (142,  68, 173),  (235, 215, 255)),
    ]

    badge_w = 108
    badge_h = 110
    gap = 14
    total_w = len(letters) * badge_w + (len(letters)-1) * gap
    start_x = (W - total_w) // 2
    y = 30

    try:
        big  = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 36)
        sm   = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 11)
        tiny = PIL.ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 9)
    except:
        big = sm = tiny = PIL.ImageFont.load_default()

    for i, (letter, label, dark, light) in enumerate(letters):
        x0 = start_x + i * (badge_w + gap)
        x1 = x0 + badge_w
        # Shadow
        draw.rounded_rectangle([x0+3, y+3, x1+3, y+badge_h+3], radius=16, fill=(180,180,180))
        # Badge fill
        draw.rounded_rectangle([x0, y, x1, y+badge_h], radius=16, fill=light, outline=dark, width=3)
        # Big letter
        lw = draw.textlength(letter, font=big)
        draw.text(((x0+x1)//2 - lw//2, y+12), letter, fill=dark, font=big)
        # Label
        lw2 = draw.textlength(label, font=sm)
        draw.text(((x0+x1)//2 - lw2//2, y+62), label, fill=dark, font=sm)
        # Separator line
        draw.line([(x0+14, y+78), (x1-14, y+78)], fill=dark, width=1)
        # Arrow connector (except last)
        if i < len(letters) - 1:
            ax = x1 + 2
            ay = y + badge_h // 2
            draw.line([(ax, ay), (ax+gap-2, ay)], fill=(120,120,120), width=2)
            draw.polygon([(ax+gap-3, ay-4), (ax+gap+3, ay), (ax+gap-3, ay+4)], fill=(120,120,120))

    # Title
    try:
        title_font = PIL.ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 16)
    except:
        title_font = PIL.ImageFont.load_default()
    title = "RICEPOT — Structured Prompt Engineering Framework"
    tw = draw.textlength(title, font=title_font)
    draw.text((W//2 - tw//2, 6), title, fill=(13,27,42), font=title_font)

    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


# ─── Document Builder ─────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    # ── Cover Page ─────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("QA AUTOMATION PORTFOLIO")
    r.bold = True
    r.font.size = Pt(32)
    r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Playwright · TypeScript · Jira MCP · AI-Driven Test Generation")
    r2.font.size = Pt(14)
    r2.font.color.rgb = TEAL

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Anand Soni  |  Senior QA Engineer  |  2026")
    r3.font.size = Pt(12)
    r3.font.color.rgb = MID_GRAY

    add_divider(doc)

    # ── Executive Summary ──────────────────────────────────────────────────────
    add_heading(doc, "Executive Summary", level=1, size=18)
    add_body(doc,
        "End-to-end test automation ecosystem built from scratch: Playwright + TypeScript with "
        "an 8-layer POM architecture, Jira integration via MCP (Model Context Protocol), "
        "AI-driven test generation and execution skills, Allure + Playwright HTML dual reporting, "
        "and two original QA engineering frameworks: Anti-Hallucination Rules and Auto-Fix Protocol. "
        "Covers 103 test cases across 5 Jira Epics with full automated lifecycle management.")

    # Stats table
    tbl = doc.add_table(rows=1, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = tbl.rows[0].cells
    stats = [("103", "Total Tests"), ("5", "Jira Epics"), ("98%", "Pass Rate"),
             ("8", "Automation Layers"), ("0", "Manual Status Updates")]
    for i, (num, label) in enumerate(stats):
        set_cell_bg(hdr_cells[i], NAVY)
        hdr_cells[i].paragraphs[0].clear()
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        r = hdr_cells[i].paragraphs[0].add_run(f"{num}\n")
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = GOLD
        r2 = hdr_cells[i].paragraphs[0].add_run(label)
        r2.font.size = Pt(9)
        r2.font.color.rgb = WHITE
    set_cell_borders(tbl)
    doc.add_paragraph()

    # ── Test Metrics Diagram ──────────────────────────────────────────────────
    add_heading(doc, "Test Results at a Glance", level=1, size=16)
    metrics_buf = make_test_metrics()
    doc.add_picture(metrics_buf, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_divider(doc)

    # ── Architecture ──────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Framework Architecture", level=1, size=18)
    add_body(doc,
        "Built on the 8-Layer POM pattern. Each layer has a single responsibility. "
        "Tests never construct page objects directly — all injected via fixture-based DI. "
        "Swapping an implementation touches one file, not every test.")

    arch_buf = make_architecture_diagram()
    doc.add_picture(arch_buf, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # DI code block
    add_heading(doc, "Fixture-Based Dependency Injection", level=2, size=13, color=TEAL)
    add_body(doc, "Single fixture file injects all page objects. Tests receive ready-to-use objects — "
             "zero construction logic in tests:", size=10)
    add_code_block(doc,
        "// src/fixtures/test-fixtures.ts\n"
        "export const test = base.extend<PageFixtures>({\n"
        "  registrationPage: async ({ page }, use) => {\n"
        "    const registrationPage = new RegistrationPage(page);\n"
        "    await use(registrationPage);  // DI: inject, use, teardown\n"
        "  },\n"
        "  blinkitLoginPage: async ({ page }, use) => {\n"
        "    const blinkitLoginPage = new BlinkitLoginPage(page);\n"
        "    await use(blinkitLoginPage);\n"
        "  },\n"
        "  // + loginPage, forgotPasswordPage\n"
        "});\n\n"
        "// Tests use it:\n"
        "test('REG-001: Verify success toast', async ({ registrationPage }) => {\n"
        "  await registrationPage.fillValidForm();\n"
        "  await registrationPage.submit();\n"
        "  await expect(registrationPage.toast).toHaveText('✅ Account created successfully!');\n"
        "});"
    )

    add_heading(doc, "Multi-Reporter Configuration", level=2, size=13, color=TEAL)
    add_code_block(doc,
        "// playwright.config.ts\n"
        "reporter: [\n"
        "  ['html',              { outputFolder: 'playwright-report', open: 'never' }],\n"
        "  ['json',              { outputFile:   'test-results/results.json' }],\n"
        "  ['list'],             // CI stdout\n"
        "  ['allure-playwright', { outputFolder: 'allure-results' }],\n"
        "],"
    )
    add_divider(doc)

    # ── AI-Driven Pipeline ────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "AI-Driven QA Pipeline", level=1, size=18)
    add_body(doc,
        "Three custom skills built as reusable Claude Code commands. The pipeline goes from "
        "Jira Epic to executed tests to updated Jira status — fully automated, zero manual steps.")

    pipeline_buf = make_pipeline_diagram()
    doc.add_picture(pipeline_buf, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Skills table
    skills_tbl = doc.add_table(rows=4, cols=3)
    skills_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Skill", "Invocation", "What It Does"]
    data_rows = [
        ["/test-case-creation",  "/test-case-creation SCRUM-142 output jira SCRUM",
         "Fetches Epic from Jira, analyzes live DOM, generates 33 test cases with requirements-traced assertions. Creates directly in Jira."],
        ["/test-case-execution", "/test-case-execution SCRUM-142",
         "Runs all child tests in headed mode, parses JSON results, updates Jira status (Done/Blocked), files bugs for app defects, archives screenshots."],
        ["/test-plan",           "/test-plan VWO-105 create docx ./output/",
         "Fetches Jira ticket, generates 14-section test plan as .md + .docx. Detects UI vs API target. Zero manual writing."],
    ]

    # Header row
    for i, h in enumerate(headers):
        cell = skills_tbl.rows[0].cells[i]
        set_cell_bg(cell, TEAL)
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_i, row_data in enumerate(data_rows):
        row = skills_tbl.rows[row_i + 1]
        bg = LIGHT_GRAY if row_i % 2 == 0 else WHITE
        for col_i, text in enumerate(row_data):
            cell = row.cells[col_i]
            set_cell_bg(cell, bg)
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(text)
            r.font.size = Pt(9.5)
            if col_i == 0:
                r.bold = True
                r.font.color.rgb = NAVY
            else:
                r.font.color.rgb = DARK_GRAY
    set_cell_borders(skills_tbl)
    add_divider(doc)

    # ── Jira MCP ──────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Jira MCP Integration — Zero-Touch Issue Management", level=1, size=18)
    add_body(doc,
        "MCP (Model Context Protocol) provides direct programmatic Jira access via OAuth. "
        "No API tokens, no copy-paste, no manual status updates. Every test execution "
        "automatically manages the full Jira lifecycle.")

    lifecycle_buf = make_jira_lifecycle()
    doc.add_picture(lifecycle_buf, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # MCP tools table
    mcp_tbl = doc.add_table(rows=9, cols=2)
    mcp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    mcp_headers = ["MCP Tool", "Used For"]
    mcp_data = [
        ["searchJiraIssuesUsingJql",    "Fetch all 33 child test cases under Epic: parent = SCRUM-142"],
        ["getJiraIssue",                "Get test case details, acceptance criteria, source file path"],
        ["transitionJiraIssue",         "To Do → In Progress → Done / Blocked (auto after test result)"],
        ["addCommentToJiraIssue",       "Execution comment: duration, retry, browser, screenshot path"],
        ["createJiraIssue",             "File Jira bug automatically when app defect detected"],
        ["createIssueLink",             "Link bug → blocked test case (Blocks relationship)"],
        ["editJiraIssue",               "Update test description with fix details, set Epic parent"],
        ["getTransitionsForJiraIssue",  "Discover valid transition IDs before transitioning status"],
    ]
    # Header
    for i, h in enumerate(mcp_headers):
        cell = mcp_tbl.rows[0].cells[i]
        set_cell_bg(cell, NAVY)
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
    for row_i, (tool, desc) in enumerate(mcp_data):
        row = mcp_tbl.rows[row_i + 1]
        bg = LIGHT_GRAY if row_i % 2 == 0 else WHITE
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], bg)
        row.cells[0].paragraphs[0].clear()
        r1 = row.cells[0].paragraphs[0].add_run(tool)
        r1.font.name = 'Courier New'
        r1.font.size = Pt(9)
        r1.font.color.rgb = TEAL
        r1.bold = True
        row.cells[1].paragraphs[0].clear()
        r2 = row.cells[1].paragraphs[0].add_run(desc)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = DARK_GRAY
    set_cell_borders(mcp_tbl)
    add_divider(doc)

    # ── Anti-Hallucination Rules ───────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Anti-Hallucination Rules Framework", level=1, size=18)
    add_body(doc,
        "Original QA engineering framework. Solves a critical problem: tests that pass because "
        "they validate wrong behavior, not because the feature works. UI-driven test creation "
        "rubber-stamps broken implementations. Requirements-driven test creation catches them.")

    ah_buf = make_ah_rules_diagram()
    doc.add_picture(ah_buf, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Rules table
    rules_tbl = doc.add_table(rows=6, cols=3)
    rules_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    rules_data = [
        ("Rule #", "Rule", "Why It Matters"),
        ("17", "Headed mode FIRST for 'element not found'",
         "Never guess selectors. 30-second investigation > 1-hour guessing."),
        ("18", "pressSequentially() for constraint tests",
         "fill() bypasses maxlength/pattern. Silent false pass. pressSequentially() simulates real keystrokes."),
        ("19", "Requirements drive assertions, UI drives locators",
         "UI typo → test still asserts correct text → test FAILS → bug caught. That's the job."),
        ("20", "Classify by actual failure mechanism",
         "XSS test failing ≠ XSS bug. Root cause may be upstream email validation. Trace first, classify after."),
        ("13", "Verify URLs and navigation paths",
         "Never assume /login, /register, /auth. WebFetch and headed mode reveal real paths."),
    ]
    hdr = True
    for row_i, (r1, r2, r3) in enumerate(rules_data):
        row = rules_tbl.rows[row_i]
        if hdr:
            for i, text in enumerate([r1, r2, r3]):
                set_cell_bg(row.cells[i], RGBColor(0x0D, 0x1B, 0x2A))
                row.cells[i].paragraphs[0].clear()
                rn = row.cells[i].paragraphs[0].add_run(text)
                rn.bold = True
                rn.font.color.rgb = WHITE
                rn.font.size = Pt(10)
            hdr = False
        else:
            bg = LIGHT_GRAY if row_i % 2 == 1 else WHITE
            set_cell_bg(row.cells[0], RGBColor(0x00, 0x87, 0x8A))
            set_cell_bg(row.cells[1], bg)
            set_cell_bg(row.cells[2], bg)
            row.cells[0].paragraphs[0].clear()
            rn = row.cells[0].paragraphs[0].add_run(f"Rule {r1}")
            rn.bold = True
            rn.font.color.rgb = WHITE
            rn.font.size = Pt(10)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[1].paragraphs[0].clear()
            rn2 = row.cells[1].paragraphs[0].add_run(r2)
            rn2.bold = True
            rn2.font.size = Pt(9.5)
            rn2.font.color.rgb = NAVY
            row.cells[2].paragraphs[0].clear()
            rn3 = row.cells[2].paragraphs[0].add_run(r3)
            rn3.font.size = Pt(9.5)
            rn3.font.color.rgb = DARK_GRAY
    set_cell_borders(rules_tbl)

    # Code example
    doc.add_paragraph()
    add_heading(doc, "Rule 18 in Code — pressSequentially() vs fill()", level=2, size=12, color=TEAL)
    add_code_block(doc,
        "// WRONG — fill() bypasses maxlength HTML attribute → FALSE PASS\n"
        "await phoneInput.fill('12345678901234');\n"
        "const v = await phoneInput.inputValue();\n"
        "expect(v.length).toBeLessThanOrEqual(10);  // ← always passes even with 14-char value\n\n"
        "// RIGHT — pressSequentially() simulates real keystrokes, browser enforces maxlength\n"
        "await phoneInput.pressSequentially('12345678901234');\n"
        "const v = await phoneInput.inputValue();\n"
        "expect(v.length).toBeLessThanOrEqual(10);  // ← REAL constraint verified"
    )
    add_divider(doc)

    # ── Auto-Fix Protocol ─────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Auto-Fix Protocol", level=1, size=18)
    add_body(doc,
        "Original incident response framework. When a test fails, investigation is systematic "
        "and autonomous — not manual. The critical gate: is this a test bug or an application bug? "
        "Wrong answer here means either broken tests staying green or valid bugs being deleted.")

    autofix_buf = make_autofix_diagram()
    doc.add_picture(autofix_buf, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Classification table
    class_tbl = doc.add_table(rows=5, cols=3)
    class_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    class_data = [
        ("Symptom", "Classification", "Action"),
        ("Element not found", "Test bug — selector wrong", "Headed mode investigation → fix selector"),
        ("Element found, click ignored", "App bug — button broken", "File Jira bug → mark test Blocked"),
        ("No error shown on invalid input", "App bug — validation missing", "File bug → DO NOT modify test assertion"),
        ("Passes first run, fails on retry", "Test bug — timing/race condition", "Fix wait strategy, not timeout value"),
    ]
    hdr = True
    for row_i, (c1, c2, c3) in enumerate(class_data):
        row = class_tbl.rows[row_i]
        if hdr:
            for i, text in enumerate([c1, c2, c3]):
                set_cell_bg(row.cells[i], DARK_GRAY)
                row.cells[i].paragraphs[0].clear()
                rn = row.cells[i].paragraphs[0].add_run(text)
                rn.bold = True
                rn.font.color.rgb = WHITE
                rn.font.size = Pt(10)
            hdr = False
        else:
            bg = LIGHT_GRAY if row_i % 2 == 1 else WHITE
            icon_color = GREEN if "App bug" not in c2 else RED
            for ci, (text, color) in enumerate([(c1, DARK_GRAY), (c2, icon_color), (c3, DARK_GRAY)]):
                set_cell_bg(row.cells[ci], bg)
                row.cells[ci].paragraphs[0].clear()
                rn = row.cells[ci].paragraphs[0].add_run(text)
                rn.font.size = Pt(9.5)
                rn.font.color.rgb = color
                if ci == 1:
                    rn.bold = True
    set_cell_borders(class_tbl)

    add_body(doc, "Critical rule: NEVER modify a test assertion to make a failing test pass "
             "if the test correctly catches an application bug. The test is evidence. Changing it destroys the evidence.",
             bold=True, color=RED, size=10)
    add_divider(doc)

    # ── Reporting ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Allure Reporting + Screenshot Archiving", level=1, size=18)
    add_body(doc,
        "Dual reporting pipeline: Allure for stakeholder dashboards (pass/fail trends, flaky detection, "
        "suite breakdown), Playwright HTML for developer debugging. Screenshots archived before every "
        "run — no failure evidence ever lost.")

    report_tbl = doc.add_table(rows=2, cols=2)
    report_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    left = report_tbl.rows[0].cells[0]
    right = report_tbl.rows[0].cells[1]
    # Merge second row
    report_tbl.rows[1].cells[0].merge(report_tbl.rows[1].cells[1])

    set_cell_bg(left, LIGHT_GRAY)
    set_cell_bg(right, LIGHT_GRAY)

    left.paragraphs[0].clear()
    rl = left.paragraphs[0].add_run("Allure Dashboard Shows")
    rl.bold = True; rl.font.size = Pt(11); rl.font.color.rgb = TEAL
    for item in ["Pass/fail % breakdown", "Suite grouping by spec file",
                 "Per-test timeline + duration", "Retry count (flaky detection)",
                 "Trend charts across runs", "Screenshot on failure"]:
        p = left.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"  •  {item}")
        r.font.size = Pt(9.5)
        r.font.color.rgb = DARK_GRAY

    right.paragraphs[0].clear()
    rr = right.paragraphs[0].add_run("Screenshot Archiving Workflow")
    rr.bold = True; rr.font.size = Pt(11); rr.font.color.rgb = TEAL
    for item in ["globalSetup archives test-results/ before every run",
                 "Playwright clears folder → archiver preserves it",
                 "Organised path: screenshots/{EpicKey}/{IssueKey}_{TestID}_{slug}_{PASS|FAIL}.png",
                 "Example: SCRUM-142/SCRUM-143_REG-001_verify-toast_PASS.png",
                 "Raw archive: screenshots-archive/{timestamp}/ (fallback)"]:
        p = right.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"  •  {item}")
        r.font.size = Pt(9.5)
        r.font.color.rgb = DARK_GRAY

    set_cell_borders(report_tbl)
    doc.add_paragraph()
    add_code_block(doc,
        "# Generate Allure report from results\n"
        "npx allure generate allure-results --clean -o allure-report\n\n"
        "# Serve in background (no CMD window)\n"
        "python -m http.server 5001 --directory allure-report &\n\n"
        "# Screenshot organised naming\n"
        "screenshots/SCRUM-142/SCRUM-152_REG-010_verify-invalid-email_FAIL.png"
    )
    add_divider(doc)

    # ── Test Matrix ───────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Full Test Coverage Matrix", level=1, size=18)

    matrix_tbl = doc.add_table(rows=7, cols=5)
    matrix_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    matrix_headers = ["Feature", "Epic", "Test IDs", "Count", "Status"]
    matrix_data = [
        ("Facebook Login",      "SCRUM-68",  "TC-001 – TC-017",    "17", "All Done ✅"),
        ("Facebook Registration","SCRUM-86", "REG-001 – REG-019",  "19", "All Done ✅"),
        ("Facebook Forgot Pwd", "SCRUM-85",  "FP-001 – FP-015",    "15", "All Done ✅"),
        ("Blinkit Login",       "SCRUM-121", "BL-001 – BL-019",    "19", "18 Done ✅  1 Blocked 🔴"),
        ("Registration Demo",   "SCRUM-142", "REG-001 – REG-033",  "33", "27 Done ✅  6 Blocked 🔴"),
        ("TOTAL",               "5 Epics",   "",                   "103","101 Done  2 Blocked"),
    ]
    for i, h in enumerate(matrix_headers):
        cell = matrix_tbl.rows[0].cells[i]
        set_cell_bg(cell, NAVY)
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_i, row_data in enumerate(matrix_data):
        row = matrix_tbl.rows[row_i + 1]
        is_total = row_data[0] == "TOTAL"
        bg = RGBColor(0xFF, 0xF9, 0xE0) if is_total else (LIGHT_GRAY if row_i % 2 == 0 else WHITE)
        for ci, text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(text)
            r.font.size = Pt(9.5)
            r.bold = is_total
            if ci == 4 and "Blocked" in text:
                r.font.color.rgb = RED
            elif ci == 4:
                r.font.color.rgb = GREEN if "Done" in text else DARK_GRAY
            else:
                r.font.color.rgb = NAVY if is_total else DARK_GRAY
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci in [1, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
    set_cell_borders(matrix_tbl)

    add_body(doc,
        "Blocked = test correctly catching an app defect. Not a test failure — a QA finding. "
        "Status stays Blocked until the bug (SCRUM-176, SCRUM-177) is fixed in the application.",
        italic=True, size=9.5, color=MID_GRAY)
    add_divider(doc)

    # ── Intentional Bugs ──────────────────────────────────────────────────────
    add_heading(doc, "Intentional Bug Design — Teaching Tool", level=1, size=16)
    add_body(doc,
        "Two bugs embedded in registration-demo.html to demonstrate realistic failure scenarios "
        "and proper QA classification workflow:")

    bugs_tbl = doc.add_table(rows=3, cols=4)
    bugs_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    bugs_headers = ["Bug ID", "Function", "Broken Behavior", "Affected Tests"]
    bugs_data = [
        ("BUG-A\nSCRUM-176", "validateEmail()",
         "Only checks @ presence.\nAccepts: notanemail@, <script>@x.com",
         "REG-009, REG-010,\nREG-011, REG-032 (XSS)"),
        ("BUG-B\nSCRUM-177", "validatePassword()",
         "Only checks length >= 8.\nAccepts: password1 (no special/upper)",
         "REG-020, REG-021"),
    ]
    for i, h in enumerate(bugs_headers):
        cell = bugs_tbl.rows[0].cells[i]
        set_cell_bg(cell, RGBColor(0xC0, 0x39, 0x2B))
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)
    for row_i, row_data in enumerate(bugs_data):
        row = bugs_tbl.rows[row_i + 1]
        set_cell_bg(row.cells[0], RGBColor(0xFF, 0xEB, 0xEB))
        set_cell_bg(row.cells[1], LIGHT_GRAY)
        set_cell_bg(row.cells[2], WHITE)
        set_cell_bg(row.cells[3], LIGHT_GRAY)
        for ci, text in enumerate(row_data):
            row.cells[ci].paragraphs[0].clear()
            r = row.cells[ci].paragraphs[0].add_run(text)
            r.font.size = Pt(9.5)
            r.font.color.rgb = RED if ci == 0 else DARK_GRAY
            r.bold = (ci == 0)
    set_cell_borders(bugs_tbl)
    add_divider(doc)

    # ── RICEPOT Framework ─────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "RICEPOT Framework — Prompt Engineering for QA", level=1, size=18)
    add_body(doc,
        "RICEPOT is a structured prompt engineering framework for crafting precise AI interactions. "
        "Used to drive all three custom skills (/test-case-creation, /test-case-execution, /test-plan) "
        "with deterministic, production-quality output. Each letter = one layer of the prompt.")

    ricepot_buf = make_ricepot_diagram()
    doc.add_picture(ricepot_buf, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # RICEPOT table
    ricepot_tbl = doc.add_table(rows=8, cols=3)
    ricepot_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    ricepot_headers = ["Letter", "Component", "Purpose + Example"]
    ricepot_data = [
        ("R", "Role",
         "Define persona AI adopts.\n"
         "Example: \"Act as a QA Functional Tester with 15 years of experience, expert in functional testing.\""),
        ("I", "Instructions",
         "Tell AI exactly what to do — no ambiguity.\n"
         "Example: \"Write complete test cases for functional + non-functional scenarios based on PRD. Use Jira format.\""),
        ("C", "Context",
         "Background: the why and where.\n"
         "Example: \"Salesforce Login Page used by enterprise sales teams. Auth via SSO. No guest access.\""),
        ("E", "Example",
         "Sample output format to guide style.\n"
         "Example: Show one test case row: Test ID, Summary, Preconditions, Steps, Expected Result, Priority."),
        ("P", "Parameters",
         "Quality + accuracy constraints.\n"
         "Example: Production-level quality. No placeholder data. Each test independently executable."),
        ("O", "Output",
         "Exact artifacts to produce — nothing else.\n"
         "Example: \"Output only the test case table. No explanations, no preamble, no commentary.\""),
        ("T", "Tone",
         "Communication style.\n"
         "Example: Technical, precise, terse. No filler language."),
    ]
    letter_colors = [
        RGBColor(0xC0, 0x39, 0x2B),  # R — red
        RGBColor(0xE6, 0x7E, 0x22),  # I — orange
        RGBColor(0xF3, 0x9C, 0x12),  # C — gold
        RGBColor(0x27, 0xAE, 0x60),  # E — green
        RGBColor(0x00, 0x87, 0x8A),  # P — teal
        RGBColor(0x29, 0x80, 0xB9),  # O — blue
        RGBColor(0x8E, 0x44, 0xAD),  # T — purple
    ]
    # Header row
    for i, h in enumerate(ricepot_headers):
        cell = ricepot_tbl.rows[0].cells[i]
        set_cell_bg(cell, NAVY)
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_i, ((letter, component, purpose), lc) in enumerate(zip(ricepot_data, letter_colors)):
        row = ricepot_tbl.rows[row_i + 1]
        light_lc = RGBColor(
            min(255, lc[0] + 160),
            min(255, lc[1] + 160),
            min(255, lc[2] + 160),
        )
        set_cell_bg(row.cells[0], lc)
        set_cell_bg(row.cells[1], light_lc)
        bg = LIGHT_GRAY if row_i % 2 == 0 else WHITE
        set_cell_bg(row.cells[2], bg)

        row.cells[0].paragraphs[0].clear()
        rl = row.cells[0].paragraphs[0].add_run(letter)
        rl.bold = True; rl.font.size = Pt(22); rl.font.color.rgb = WHITE
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row.cells[1].paragraphs[0].clear()
        rc = row.cells[1].paragraphs[0].add_run(component)
        rc.bold = True; rc.font.size = Pt(11); rc.font.color.rgb = NAVY

        row.cells[2].paragraphs[0].clear()
        lines = purpose.split('\n')
        rp = row.cells[2].paragraphs[0].add_run(lines[0])
        rp.font.size = Pt(9.5); rp.font.color.rgb = DARK_GRAY; rp.bold = True
        if len(lines) > 1:
            p2 = row.cells[2].add_paragraph()
            p2.paragraph_format.space_before = Pt(2)
            ri = p2.add_run(lines[1])
            ri.font.size = Pt(9); ri.font.color.rgb = MID_GRAY; ri.italic = True

    set_cell_borders(ricepot_tbl)
    doc.add_paragraph()

    # Real usage example
    add_heading(doc, "RICEPOT in Action — /test-case-creation Skill Prompt", level=2, size=12, color=TEAL)
    add_body(doc, "How RICEPOT structures the prompt that drives /test-case-creation:", size=10)
    add_code_block(doc,
        "Role:         Senior QA Engineer, expert in Playwright + Jira\n"
        "Instructions: Generate test cases from Epic SCRUM-142 acceptance criteria.\n"
        "              Fetch Epic via MCP. Analyze DOM via fetch-local-page.js.\n"
        "              Create 33 Jira issues under SCRUM-142 with exact expected results.\n"
        "Context:      Registration demo at localhost:7000/registration-demo.html.\n"
        "              Known intentional bugs: BUG-A (email), BUG-B (password).\n"
        "              Anti-Hallucination Rule 19: assertions from Epic, locators from DOM.\n"
        "Example:      Test ID: REG-001 | Expected: '✅ Account created successfully!'\n"
        "              Source column: 'Epic SCRUM-142' (not UI Observed)\n"
        "Parameters:   No invented behavior. Every assertion traceable to Epic.\n"
        "              Flag requirement gaps. Security tests always included.\n"
        "Output:       33 Jira issues created, linked to Epic, Jira keys returned.\n"
        "Tone:         Technical, terse, zero filler."
    )

    add_heading(doc, "Anti-Hallucination Rules Built Into RICEPOT", level=2, size=12, color=TEAL)
    add_body(doc, "RICEPOT's Parameters section always embeds AH constraints:", size=10)
    ah_tbl = doc.add_table(rows=5, cols=2)
    ah_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    ah_constraints = [
        ("AH Constraint", "RICEPOT Section"),
        ("Do not assume default or typical system behaviour", "Parameters (P)"),
        ("If information missing → 'Insufficient information to determine'", "Parameters (P)"),
        ("If detail inferred → label as 'Inferred (low confidence)'", "Parameters (P)"),
        ("Output must be deterministic, repeatable, follow exact format", "Output (O) + Tone (T)"),
    ]
    for row_i, (c1, c2) in enumerate(ah_constraints):
        row = ah_tbl.rows[row_i]
        if row_i == 0:
            set_cell_bg(row.cells[0], TEAL)
            set_cell_bg(row.cells[1], TEAL)
            for ci, text in enumerate([c1, c2]):
                row.cells[ci].paragraphs[0].clear()
                r = row.cells[ci].paragraphs[0].add_run(text)
                r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)
        else:
            bg = LIGHT_GRAY if row_i % 2 == 1 else WHITE
            set_cell_bg(row.cells[0], bg)
            set_cell_bg(row.cells[1], bg)
            row.cells[0].paragraphs[0].clear()
            r1 = row.cells[0].paragraphs[0].add_run(c1)
            r1.font.size = Pt(9.5); r1.font.color.rgb = DARK_GRAY
            row.cells[1].paragraphs[0].clear()
            r2 = row.cells[1].paragraphs[0].add_run(c2)
            r2.font.size = Pt(9.5); r2.font.color.rgb = TEAL; r2.bold = True
    set_cell_borders(ah_tbl)
    add_divider(doc)

    # ── Accuracy Analysis ─────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Framework Impact — Accuracy Analysis", level=1, size=18)
    add_body(doc,
        "Evidence-based accuracy numbers from real work on this project. "
        "Numbers derived from actual failures caught, selectors hallucinated, "
        "and bugs misclassified — before and after applying the frameworks.")

    acc_buf = make_accuracy_chart()
    doc.add_picture(acc_buf, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Detailed breakdown table
    acc_tbl = doc.add_table(rows=9, cols=4)
    acc_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    acc_headers = ["Area", "Without Frameworks", "With Frameworks", "Key Rule / Reason"]
    acc_data = [
        ("SPA Selector Accuracy",
         "17–45%\n(Facebook: 5/6 hallucinated)",
         "95–98%",
         "Rule 17: headed mode\nbefore any selector used"),
        ("Variable Naming Accuracy",
         "20–60%\n(searchButton, signUpButton, etc.)",
         "~98%",
         "Auto-Fix Rule 13:\nfull semantic rename required"),
        ("Test Assertion Correctness",
         "40–55%\n(rubber-stamps broken UI)",
         "~97%",
         "Rule 19: Epic = assertions\nDOM = locators only"),
        ("Bug Classification",
         "30–40%\n(XSS test fail → filed as XSS bug)",
         "~95%",
         "Rule 20: trace to first\nbroken gate in input chain"),
        ("Constraint Test Accuracy",
         "0%\n(fill() bypasses maxlength always)",
         "~98%",
         "Rule 18: pressSequentially()\nsimulates real keystrokes"),
        ("URL/Navigation Accuracy",
         "25–35%\n(/login, /forgot all assumed wrong)",
         "~97%",
         "Rule 13: verify via\nWebFetch / headed mode"),
        ("Jira Status Accuracy",
         "60–70%\n(manual → forgotten / wrong status)",
         "~99%",
         "MCP automation:\nzero manual updates"),
        ("Overall Defect Detection",
         "~0% effective\n(passing tests, broken app)",
         "~95%+",
         "Requirements-driven\nassertions catch real bugs"),
    ]

    # Header
    for i, h in enumerate(acc_headers):
        cell = acc_tbl.rows[0].cells[i]
        set_cell_bg(cell, NAVY)
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_i, (area, without, with_fw, rule) in enumerate(acc_data):
        row = acc_tbl.rows[row_i + 1]
        bg = LIGHT_GRAY if row_i % 2 == 0 else WHITE
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], RGBColor(0xFF, 0xEB, 0xEB))
        set_cell_bg(row.cells[2], RGBColor(0xE8, 0xF8, 0xEE))
        set_cell_bg(row.cells[3], bg)

        row.cells[0].paragraphs[0].clear()
        r0 = row.cells[0].paragraphs[0].add_run(area)
        r0.bold = True; r0.font.size = Pt(9.5); r0.font.color.rgb = NAVY

        row.cells[1].paragraphs[0].clear()
        lines_w = without.split('\n')
        r1a = row.cells[1].paragraphs[0].add_run(lines_w[0])
        r1a.bold = True; r1a.font.size = Pt(10); r1a.font.color.rgb = RED
        if len(lines_w) > 1:
            p2 = row.cells[1].add_paragraph()
            p2.paragraph_format.space_before = Pt(1)
            r1b = p2.add_run(lines_w[1])
            r1b.font.size = Pt(8.5); r1b.font.color.rgb = MID_GRAY; r1b.italic = True

        row.cells[2].paragraphs[0].clear()
        r2 = row.cells[2].paragraphs[0].add_run(with_fw)
        r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = GREEN
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row.cells[3].paragraphs[0].clear()
        lines_r = rule.split('\n')
        r3a = row.cells[3].paragraphs[0].add_run(lines_r[0])
        r3a.bold = True; r3a.font.size = Pt(9); r3a.font.color.rgb = TEAL
        if len(lines_r) > 1:
            p3 = row.cells[3].add_paragraph()
            p3.paragraph_format.space_before = Pt(1)
            r3b = p3.add_run(lines_r[1])
            r3b.font.size = Pt(8.5); r3b.font.color.rgb = DARK_GRAY

    set_cell_borders(acc_tbl)
    doc.add_paragraph()

    # The most dangerous gap callout
    add_heading(doc, "The Most Dangerous Gap", level=2, size=13, color=RED)
    add_body(doc,
        "Assertion accuracy without Rule 19 sits at 40–55% — but this is the worst failure mode "
        "because the tests go GREEN while catching nothing:", size=10)
    add_code_block(doc,
        "// WITHOUT Rule 19 — UI-driven assertion (rubber-stamp)\n"
        "// UI shows typo 'Sinup' → test asserts 'Sinup' → PASS ✅\n"
        "// Dev shipped typo. QA said it's fine. Release went out broken.\n"
        "await expect(page.getByText('Sinup')).toBeVisible();\n\n"
        "// WITH Rule 19 — Requirements-driven assertion\n"
        "// Epic says 'Sign Up' → test asserts 'Sign Up' → FAIL ❌\n"
        "// Bug caught before release. That IS the job.\n"
        "await expect(page.getByText('Sign Up')).toBeVisible();"
    )
    add_body(doc,
        "A 95% pass rate WITHOUT Rule 19 may have 0% real defect detection. "
        "An 85% pass rate WITH Rule 19 is genuinely finding bugs. "
        "Pass rate and quality are NOT the same metric.",
        bold=True, color=RED, size=10)
    add_divider(doc)

    # ── Correction Cycle Time ─────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Real Cost of Skipping the Frameworks", level=1, size=18)
    add_body(doc,
        "Both approaches use AI. The difference: WITHOUT frameworks, AI makes plausible mistakes "
        "that look correct. User detects them, prompts corrections, AI fixes, user re-checks. "
        "Each cycle burns 15–60 min. WITH frameworks baked in, most mistakes are prevented before output. "
        "Below is the real correction-cycle cost per Epic (30 tests) measured from this project.")

    cyc_buf = make_correction_cycle_chart()
    doc.add_picture(cyc_buf, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Correction cycle table
    cyc_tbl = doc.add_table(rows=8, cols=4)
    cyc_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cyc_headers = ["Mistake Type", "AI Without Frameworks", "Correction Cost", "WITH Frameworks"]
    cyc_data = [
        ("Wrong selectors\n(SPA hallucination)",
         "5–6 selectors wrong per page\nUser: 'this selector fails'\nAI: guesses new one (wrong again)\n→ 3–5 re-prompt cycles",
         "45–90 min\nper page object",
         "Rule 17: headed mode first\n0 re-prompts\n~5 min verified"),
        ("Wrong variable names\n(searchButton etc.)",
         "Names look correct but wrong\nFound only when test runs\nFull rename across 4 files\n→ 2–3 correction cycles",
         "30–60 min\nper POM class",
         "Rule 13: full semantic rename\non first detection\n~5 min"),
        ("UI-observed assertions\n(rubber-stamp tests)",
         "Tests pass, bugs slip through\nFound in prod or user testing\nAll assertions need rewrite\n→ entire test suite rewrite",
         "4–8 hours\nrewrite sprint",
         "Rule 19: Epic-first assertions\nCorrect from day 1\n~0 rework"),
        ("Wrong bug classification\n(XSS vs email bug)",
         "Bug filed wrong\nDev team investigates wrong area\nReopened, re-investigated\n→ 2–3 sprint cycles wasted",
         "2–4 hours\nper misclassified bug",
         "Rule 20: trace to first gate\nCorrect classification\n~10 min"),
        ("fill() for constraints\n(silent false pass)",
         "Tests green, constraint broken\nFound in prod when user types\nAll constraint tests rewrite\n→ discovered post-release",
         "Unknown — post-release\n+ hotfix cost",
         "Rule 18: pressSequentially()\nCaught in dev\n~0 rework"),
        ("Manual Jira updates\n(30 tests per Epic)",
         "Each update: open Jira, find issue,\ntransition, add comment, attach screenshot\n→ 30 × 3–4 min",
         "90–120 min\nper Epic run",
         "MCP: fully automated\n~0 min per run"),
    ]

    for i, h in enumerate(cyc_headers):
        cell = cyc_tbl.rows[0].cells[i]
        set_cell_bg(cell, DARK_GRAY)
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)

    for row_i, (mtype, without, cost, with_fw) in enumerate(cyc_data):
        row = cyc_tbl.rows[row_i + 1]
        bg = LIGHT_GRAY if row_i % 2 == 0 else WHITE
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], RGBColor(0xFF, 0xEB, 0xEB))
        set_cell_bg(row.cells[2], RGBColor(0xFF, 0xF3, 0xE0))
        set_cell_bg(row.cells[3], RGBColor(0xE8, 0xF8, 0xEE))

        # col 0: mistake type
        row.cells[0].paragraphs[0].clear()
        lines = mtype.split('\n')
        r0 = row.cells[0].paragraphs[0].add_run(lines[0])
        r0.bold = True; r0.font.size = Pt(9.5); r0.font.color.rgb = NAVY
        for l in lines[1:]:
            p = row.cells[0].add_paragraph(); p.paragraph_format.space_before = Pt(1)
            ri = p.add_run(l); ri.font.size = Pt(8.5); ri.font.color.rgb = MID_GRAY; ri.italic = True

        # col 1: without
        row.cells[1].paragraphs[0].clear()
        for li, l in enumerate(without.split('\n')):
            if li == 0:
                r1 = row.cells[1].paragraphs[0].add_run(l)
                r1.font.size = Pt(9); r1.font.color.rgb = RED
            else:
                p = row.cells[1].add_paragraph(); p.paragraph_format.space_before = Pt(1)
                ri = p.add_run(l); ri.font.size = Pt(8.5); ri.font.color.rgb = DARK_GRAY

        # col 2: cost
        row.cells[2].paragraphs[0].clear()
        lines_c = cost.split('\n')
        r2 = row.cells[2].paragraphs[0].add_run(lines_c[0])
        r2.bold = True; r2.font.size = Pt(10); r2.font.color.rgb = ORANGE
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for l in lines_c[1:]:
            p = row.cells[2].add_paragraph(); p.paragraph_format.space_before = Pt(1)
            ri = p.add_run(l); ri.font.size = Pt(8.5); ri.font.color.rgb = MID_GRAY
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # col 3: with frameworks
        row.cells[3].paragraphs[0].clear()
        for li, l in enumerate(with_fw.split('\n')):
            if li == 0:
                r3 = row.cells[3].paragraphs[0].add_run(l)
                r3.bold = True; r3.font.size = Pt(9); r3.font.color.rgb = GREEN
            else:
                p = row.cells[3].add_paragraph(); p.paragraph_format.space_before = Pt(1)
                ri = p.add_run(l); ri.font.size = Pt(8.5); ri.font.color.rgb = DARK_GRAY

    set_cell_borders(cyc_tbl)
    doc.add_paragraph()

    # Total time callout
    add_heading(doc, "Total Time Per Epic (30 Tests) — AI-Assisted Both Ways", level=2, size=13, color=NAVY)
    tot_tbl = doc.add_table(rows=4, cols=3)
    tot_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tot_headers = ["", "AI WITHOUT Frameworks", "AI WITH Frameworks"]
    tot_data = [
        ("Generation time\n(Epic + tests + POM + spec)", "30–45 min", "15–25 min"),
        ("Correction cycles\n(re-prompts, fixes, re-runs)", "6–14 hours", "0–30 min"),
        ("TOTAL per Epic", "~7–15 hours", "~20–30 min"),
    ]
    for i, h in enumerate(tot_headers):
        cell = tot_tbl.rows[0].cells[i]
        set_cell_bg(cell, NAVY)
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_i, (label, without, with_fw) in enumerate(tot_data):
        row = tot_tbl.rows[row_i + 1]
        is_total = row_i == 2
        set_cell_bg(row.cells[0], LIGHT_GRAY)
        set_cell_bg(row.cells[1], RGBColor(0xFF, 0xEB, 0xEB) if not is_total else RGBColor(0xFF, 0xCC, 0xCC))
        set_cell_bg(row.cells[2], RGBColor(0xE8, 0xF8, 0xEE) if not is_total else RGBColor(0xC8, 0xF0, 0xD8))

        row.cells[0].paragraphs[0].clear()
        lines = label.split('\n')
        r0 = row.cells[0].paragraphs[0].add_run(lines[0])
        r0.bold = is_total; r0.font.size = Pt(10 if is_total else 9.5); r0.font.color.rgb = NAVY
        if len(lines) > 1:
            p = row.cells[0].add_paragraph(); p.paragraph_format.space_before = Pt(1)
            ri = p.add_run(lines[1]); ri.font.size = Pt(8.5); ri.font.color.rgb = MID_GRAY; ri.italic = True

        row.cells[1].paragraphs[0].clear()
        r1 = row.cells[1].paragraphs[0].add_run(without)
        r1.bold = is_total; r1.font.size = Pt(12 if is_total else 10)
        r1.font.color.rgb = RED
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row.cells[2].paragraphs[0].clear()
        r2 = row.cells[2].paragraphs[0].add_run(with_fw)
        r2.bold = is_total; r2.font.size = Pt(12 if is_total else 10)
        r2.font.color.rgb = GREEN
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_cell_borders(tot_tbl)

    add_body(doc,
        "The frameworks don't just save time — they change where time is spent. "
        "Without them: 80–90% of time is correction cycles (detecting mistakes, re-prompting, re-verifying). "
        "With them: 90%+ of time is on actual QA thinking — reviewing results, classifying bugs, improving coverage.",
        bold=False, color=DARK_GRAY, size=10)
    add_divider(doc)

    # ── Interview Q&A ─────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Interview Q&A — Prepared Answers", level=1, size=18)

    qa_pairs = [
        ("How do you handle flaky tests?",
         "Flaky tests have a root cause. My Auto-Fix Protocol classifies: timing (add proper wait), "
         "networkidle (change wait strategy), selector fragility (fix to stable ID/attribute). "
         "Allure shows retry > 0 — anything flaky gets investigated before merge. "
         "force: true and waitForTimeout are symptoms, not fixes."),
        ("How do you ensure tests don't validate broken behavior?",
         "Anti-Hallucination Rule 19: UI analysis gives locators, requirements give assertions. "
         "I never derive expected results from what the UI currently does. If the UI disagrees with "
         "the Epic — the test fails. That IS the job. The test is correct, the UI is wrong."),
        ("How do you manage selector maintenance?",
         "ID-based selectors — most stable type. Every locator verified in headed mode before use, "
         "with verification date in comment. When a selector breaks: Rule 17 — headed mode first, "
         "then fix the page object helper method, not individual tests. One fix covers all consumers."),
        ("What is your Jira integration approach?",
         "Full lifecycle via MCP: To Do → In Progress → Done or Blocked, all automated after test run. "
         "Bug filing is automatic — skill creates the Jira issue, links it to the blocked test, "
         "adds comment with classification reason. Engineers do zero manual status updates."),
        ("How do you distinguish test failures from app bugs?",
         "Auto-Fix Protocol Step 3 (Classify): run headed mode, manually test the scenario. "
         "If manual test in browser also fails → app bug. If only automated test fails → test issue. "
         "Never modify assertions to pass a failing test catching a real bug — the test is evidence."),
        ("What is your approach to test data?",
         "Dedicated test data layer (Layer 3). Test data never hardcoded in tests. "
         "Separate test data for positive/negative/edge cases. BL-018 uses SQL injection payload "
         "as test data — traced from Epic security requirement, not invented by tester."),
    ]

    for q, a in qa_pairs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        rq = p.add_run(f"Q: {q}")
        rq.bold = True
        rq.font.size = Pt(11)
        rq.font.color.rgb = NAVY

        pa = doc.add_paragraph()
        pa.paragraph_format.left_indent = Inches(0.25)
        pa.paragraph_format.space_before = Pt(2)
        pa.paragraph_format.space_after = Pt(8)
        ra = pa.add_run(f"A: {a}")
        ra.font.size = Pt(10.5)
        ra.font.color.rgb = DARK_GRAY

    add_divider(doc)

    # ── Footer ─────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Anand Soni  ·  QA Automation Engineer  ·  anandsoni2641@gmail.com  ·  2026")
    r.font.size = Pt(10)
    r.font.color.rgb = MID_GRAY
    r.italic = True

    return doc

# ─── Main ─────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    output_path = r"c:\ClaudeCodeMasterclass\QA-Automation-Portfolio-Anand-Soni.docx"
    print("Building portfolio document...")
    doc = build_doc()
    doc.save(output_path)
    print(f"Saved: {output_path}")
